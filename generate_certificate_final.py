import os
import pandas as pd
from docx import Document
from datetime import datetime

# ---------- Helpers ----------
def format_date(date):
    if pd.isna(date): return ""
    return pd.to_datetime(date).strftime("%#d %B %Y") if os.name == 'nt' else pd.to_datetime(date).strftime("%-d %B %Y")

def calculate_months(start, end):
    start = pd.to_datetime(start)
    end = pd.to_datetime(end)
    return round((end - start).days / 30)

def replace_text_preserving_format(paragraph, replacements):
    for run in paragraph.runs:
        for key, value in replacements.items():
            if key in run.text:
                run.text = run.text.replace(key, value)

# ---------- Paths ----------
DATA_FILE = "data/student_data.xlsx"
TEMPLATE_FILE = "templates/internship_template.docx"
OUTPUT_DIR = "output"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# ---------- Load Data ----------
df = pd.read_excel(DATA_FILE)
df.columns = [col.strip().upper() for col in df.columns]

# ---------- Process Each Student ----------
for idx, row in df.iterrows():
    name = row["NAME"]
    domain = row["DOMAIN"]
    start_date = format_date(row["START DATE"])
    end_date = format_date(row["END DATE"])
    issue_date = format_date(row["ISSUE DATE"])
    duration = f"{calculate_months(row['START DATE'], row['END DATE'])} month"

    # Load Word template
    doc = Document(TEMPLATE_FILE)

    # Replacements map
    replacements = {
        "{Student Name}": name,
        "{Domain Name}": domain,
        "{Start date}": start_date,
        "{End date}": end_date,
        "{Issue date}": issue_date,
        "{No of months}": duration
    }

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        replace_text_preserving_format(paragraph, replacements)

    # Replace in tables if present
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:
                for paragraph in cell.paragraphs:
                    replace_text_preserving_format(paragraph, replacements)

    # Save with student name
    filename = f"{name.replace(' ', '_')}_Internship_Certificate.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))

print("✅ Certificates successfully generated for all students!")
