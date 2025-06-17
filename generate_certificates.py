import os
import pandas as pd
from docx import Document
from datetime import datetime

# Format date for display
def format_date(date):
    if pd.isna(date):
        return ""
    return pd.to_datetime(date).strftime("%#d %B %Y") if os.name == 'nt' else pd.to_datetime(date).strftime("%-d %B %Y")

# Calculate months
def calculate_months(start, end):
    start = pd.to_datetime(start)
    end = pd.to_datetime(end)
    return round((end - start).days / 30)

# Paths
DATA_FILE = "data/student_data.xlsx"
TEMPLATE_FILE = "templates/internship_template.docx"
OUTPUT_DIR = "output"

os.makedirs(OUTPUT_DIR, exist_ok=True)

# Load Excel
df = pd.read_excel(DATA_FILE)
df.columns = [col.strip().upper() for col in df.columns]

# Loop over each student
for idx, row in df.iterrows():
    name = row["NAME"]
    domain = row["DOMAIN"]
    start_date = format_date(row["START DATE"])
    end_date = format_date(row["END DATE"])
    issue_date = format_date(row["ISSUE DATE"])
    duration_months = calculate_months(row["START DATE"], row["END DATE"])

    # Load template
    doc = Document(TEMPLATE_FILE)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        paragraph.text = paragraph.text.replace("{Student Name}", name)
        paragraph.text = paragraph.text.replace("{Domain Name}", domain)
        paragraph.text = paragraph.text.replace("{Start date}", start_date)
        paragraph.text = paragraph.text.replace("{End date}", end_date)
        paragraph.text = paragraph.text.replace("{No of months}", f"{duration_months} month{'s' if duration_months > 1 else ''}")
        paragraph.text = paragraph.text.replace("{Issue date}", issue_date)

    # Also replace in tables if any
    for table in doc.tables:
        for row_ in table.rows:
            for cell in row_.cells:
                for p in cell.paragraphs:
                    p.text = p.text.replace("{Student Name}", name)
                    p.text = p.text.replace("{Domain Name}", domain)
                    p.text = p.text.replace("{Start date}", start_date)
                    p.text = p.text.replace("{End date}", end_date)
                    p.text = p.text.replace("{No of months}", f"{duration_months} month{'s' if duration_months > 1 else ''}")
                    p.text = p.text.replace("{Issue date}", issue_date)

    # Save file
    filename = f"{name.replace(' ', '_')}_Internship_Certificate.docx"
    doc.save(os.path.join(OUTPUT_DIR, filename))

print("✅ Certificates generated for all students in 'output/' folder.")
